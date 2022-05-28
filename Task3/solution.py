from docx import Document
import os


def is_doc_contract(file_name):

    Listbolds = []
    doc = Document(file_name)

    for para in doc.paragraphs:
        for run in para.runs:
            if run.bold:
                word = run.text
                Listbolds.append(word)

    if Listbolds == []:
        return "unknown"
    else:
        for item in Listbolds:
            if item.find("Agreement") or item.find("Law") or item.find("Term"):
                return "Contract"
            else:
                return "unknown"

def is_doc_letter(file_name):
    
    doc = Document(file_name)
    is_letter = False

    for para in doc.paragraphs:      
        
        if para.text.find("Dear ") > -1 or para.text.find("Hi ") > -1:
            is_letter = True
    
    if is_letter == False:
        return "unknown"
    else:
        return "Letter"


def get_doc_type(dir_path):

    doc_type_all = {}
    doc_files_list = []

    for file_name in os.listdir(dir_path):
        if file_name.endswith(".docx"):
            doc_files_list.append(dir_path + file_name)

    for file_name in doc_files_list:

        doc_type = is_doc_contract(file_name)

        if doc_type == "unknown":
            doc_type = is_doc_letter(file_name)
            if doc_type == "unknown":
                doc_type_all[file_name] = "Contract"
            else:
                doc_type_all[file_name] = "Letter"

        else:
            doc_type_all[file_name] = "Contract"

    return doc_type_all



if __name__ == "__main__":
    
    dir_path = "test_docs/"
    doc_type_all = get_doc_type(dir_path)

    print("\nDocument Types in " + dir_path + ": ")

    for file_name, type in doc_type_all.items():
        print(file_name + " : " + type)