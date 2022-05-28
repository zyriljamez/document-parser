from docx import Document
import os
import re
import calendar

months = list(calendar.month_name)[1:]


def get_address_contract(file_name):

    doc = Document(file_name)
    address_list = []    
    address_add = False
    curr_address = ""


    for para in doc.paragraphs:

        #single line address
        match = re.search("address\sis\sat\s(.+)hereinafter", para.text)
        if match is not None:
            curr_address = match.group(1)
            address_list.append(curr_address)
            curr_address = ""
        else:
            #multiline address
            if address_add is True and para.text.find("hereinafter") == -1 and para.text.find("hereby") == -1:
                curr_address = curr_address + "\n" + para.text
            
            if para.text.find("address is") > -1:
                address_add = True
            
            if para.text.find("hereinafter") > -1 or para.text.find("hereby") > -1:
                address_add = False
                if curr_address != "":
                    address_list.append(curr_address)
                    curr_address = ""

    return address_list    


def get_address_letter(file_name):

    doc = Document(file_name)
    address_list = []
    address_add = True
    curr_address = ""
    is_letter = False

    for para in doc.paragraphs:

        if address_add is True and para.text.find("Dear ") == -1 and para.text.find("Hi ") == -1:
            date_string = False
            for month in months:
                if para.text.find(month) > -1:
                    address_list.append(curr_address)
                    curr_address = ""
                    date_string = True
                    break
            if date_string != True:    
                curr_address = curr_address + "\n" + para.text
        
        
        if para.text.find("Dear ") > -1 or para.text.find("Hi ") > -1:
            is_letter = True
            address_add = False
            if curr_address != "":
                address_list.append(curr_address)
            break
        
    if is_letter == False:
        return []
    else:
        return address_list


def get_address_list(dir_path):

    doc_files_list = []

    for file_name in os.listdir(dir_path):
        if file_name.endswith(".docx"):
            doc_files_list.append(dir_path + file_name)

    address_list_all = []

    for file_name in doc_files_list:

        address_list = get_address_contract(file_name)

        if address_list == []:
            address_list = get_address_letter(file_name)

        if address_list != []:
            for address in address_list:
                address_list_all.append(address)

    return address_list_all


if __name__ == "__main__":

    dir_path = "test_docs/"
    address_list_all = get_address_list(dir_path)

    print("\nList of addresses in " + dir_path + ": ")
    for address in address_list_all:
        print(address)
